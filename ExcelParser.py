#!/usr/bin/env python
#-*- encoding: utf-8 -*-


import os
import tkinter
import tkinter.filedialog
import tkinter.messagebox
import re
import Levenshtein
import datetime
import webbrowser

from tkinter import ttk
from decimal import Decimal
from xlrd import biffh, open_workbook
from docx import Document
from docx.oxml.ns import qn

# from MultiListBox import MultiListBox


def add2dict(dict_x, key_x, set_x):
  if key_x in dict_x:
    dict_x[key_x] = set_x
  else:
    dict_x.update({key_x: set_x})


# 正则匹配(千位分隔符可选)浮点数，用于纠正Excel中可能出现的单元格数据格式问题
def matchNumber(string):
  tmp = str(string)
  reg = re.compile(r'^[-+]?([0-9]{1,3}(,?[0-9]{3})*(\.[0-9]+)?|\.[0-9]+)$')
  if reg.match(tmp):
    return True
  return False


def getCell(sheet, index_x, index_y):
  return sheet.cell_type(index_x, index_y), sheet.cell_value(index_x, index_y)


def preTreat(string):
  string = string.replace('*', '')
  string = string.replace('△', '')
  string = string.replace(' ', '')  # 预处理，去除空格
  lq = string.find('（')
  rq = string.find('）')
  while lq != -1 and rq != -1 and lq < rq:
    tmp = string[lq:rq+1]
    string = string.replace(tmp, '')
    lq = string.find('（')
    rq = string.find('）')

  pm = string.find('、')
  while pm != -1 and pm < 3:
    tmp = string[:pm+1]
    string = string.replace(tmp, '')
    pm = string.find('、')

  return string



'''
def isFind(string):
  keywords = ['非流动资产合计', '流动资产合计', '资产合计', '资产总计', '非流动负债合计', '流动负债合计', '负债总计', '负债合计', '实收资本', '所有者权益合计', '营业收入', '营业成本', '利润总额', '净利润']
  for keyword in keywords:
    if string == keyword or str(string).find(keyword) != -1:
      return True, keyword
  return False, ''
'''


def countCell(sheet, index, i, j):
  tmp_y = j + 1
  if tmp_y + 1 in index:
    tmp_y = tmp_y + 1
  elif tmp_y + 2 in index:
    tmp_y = tmp_y + 2
  type_start, start = getCell(sheet, i, tmp_y+1)
  type_end, end = getCell(sheet, i, tmp_y)
  change = ''
  # 重复代码，待优化
  # 重复代码，待优化
  # 重复代码，待优化
  if type_start == biffh.XL_CELL_NUMBER and type_end == biffh.XL_CELL_NUMBER:
    start = Decimal(str(start))
    end = Decimal(str(end))
    change = ((end - start)/end*Decimal(100)).quantize(Decimal('0.00'))
    start = (start/Decimal(10000)).quantize(Decimal('0.00'))
    end = (end/Decimal(10000)).quantize(Decimal('0.00'))
  elif (type_start == biffh.XL_CELL_TEXT and matchNumber(start)) and (type_end == biffh.XL_CELL_TEXT and matchNumber(end)):
    start = Decimal(start)
    end = Decimal(end)
    change = ((end - start)/end*Decimal(100)).quantize(Decimal('0.00'))
    start = (start/Decimal(10000)).quantize(Decimal('0.00'))
    end = (end/Decimal(10000)).quantize(Decimal('0.00'))
  #tmp = {'年初': start, '期末': end, '变动': change}
  tmp = {'start': start, 'end': end, 'change': change}
  return tmp


'''
def countSheet(wb, sheet_name):
  sheet = wb.sheet_by_name(sheet_name)
  n_rows = sheet.nrows
  n_cols = sheet.ncols
  result_dict = {}
  start_index = []
  end_index = []
  for i in range(n_rows):
    for j in range(n_cols):
      index_x_y = sheet.cell_value(i, j)
      if index_x_y == '期末余额' or index_x_y == '本期金额':
        end_index.append(j)
      if index_x_y == '年初余额' or index_x_y == '上期金额':
        start_index.append(j)

      flag, key = isFind(index_x_y)
      if(flag):
        tmp = countCell(sheet, end_index, i, j)
        add2dict(result_dict, key, tmp)
  return result_dict
'''


class ParseExcel(object):
  def __init__(self):
    # 初始化主窗口
    self.root = tkinter.Tk()
    self.root.resizable(0, 0)
    self.center_main_window(640, 420)
    self.root.iconbitmap('favicon.ico')
    self.root.title('报表解析')

    # 声明全局变量
    self.filename = ''
    self.foldername = ''
    self.keywords = []
    self.results = {}

    # 定义Frame进行布局
    self.top = ttk.Frame()
    s = ttk.Style()
    s.configure('E.TEntry', borderwidth=10, insertwidth=1, relief=tkinter.FLAT)
    self.file_entry = ttk.Entry(self.top, width=62, style='E.TEntry')
    # self.file_button = ttk.Button(self.top, command=self.chooseFile, text='选择文件', width=10)
    self.parse_button = ttk.Button(self.top, command=self.conduct, text='选择并解析', width=10)
    self.export_button = ttk.Button(self.top, command=self.export, text='导出结果', width=10)


    self.middle = tkinter.Frame()
    # self.result_list = tkinter.Listbox(self.middle, width=90, height=20, relief=tkinter.FLAT)
    # self.result_list = MultiListBox(self.middle, (('科目', 20), ('年初（万元）', 20), ('期末（万元）', 20), ('变动（%）', 20)), height=18)
    # 因为MultiListBox的kw参数中需要height，因此height为必选参数，详情见MultiListBox.py

    self.result_list = ttk.Treeview(self.middle, show='headings', height=15, columns=('subject', 'start', 'end', 'change'))

    self.result_list.column('subject', width=158, anchor="center")
    self.result_list.column('start', width=156, anchor="center")
    self.result_list.column('end', width=156, anchor="center")
    self.result_list.column('change', width=142, anchor="center")

    self.result_list.heading('subject', text='科目')
    self.result_list.heading('start', text='年初（万元）')
    self.result_list.heading('end', text='期末（万元）')
    self.result_list.heading('change', text='变动（%）')

    # self.scroll_bar = ttk.Scrollbar(self.middle, orient=tkinter.VERTICAL, command=self.result_list.yview)

    # self.result_list.configure(yscrollcommand=self.scroll_bar.set)

    self.bottom = ttk.Frame()
    now = datetime.datetime.now().strftime('%b %d %Y %H:%M:%S')
    self.author_label = ttk.Label(self.bottom, text='Version 0.10, Code by Levan, Build on {}'.format(now))
    self.author_button = ttk.Button(self.bottom, command=self.broswer, text='?', width=2)


  # 获取屏幕的大小，用于主窗口的居中
  def get_screen_size(self):
    return self.root.winfo_screenwidth(), self.root.winfo_screenheight()


  # 获取主窗口的大小，用于子窗口的定位
  def get_window_size(self):
    return self.root.winfo_reqwidth(), self.root.winfo_reqheight()


  # 获取主窗口的位置，用于子窗口的定位
  def get_window_pos(self):
    self.root.update()
    return self.root.winfo_x(), self.root.winfo_y()


  # 主窗口中心点位置
  def main_window_pos(self):
    window_x, window_y = self.get_window_pos()
    window_w, window_h = self.get_window_size()
    # main_center_x, main_center_y = window_x+window_w/2, window_y+window_h/2
    # 为了统一使用center_window函数，这里将main_center_x、main_center_y的结果同时乘以2
    return window_x*2+window_w, window_y*2+window_h


  # 窗口居中显示
  # 居中定位函数可能会被多个组件调用，因此引入window做参数，而不是self.root将其写死
  def center_window(self, window, width, height, screenwidth, screenheight):
    size = '%dx%d+%d+%d' % (width, height, (screenwidth - width)/2, (screenheight - height)/2)
    window.geometry(size)

  # 主窗口居中
  def center_main_window(self, width, height):
    screenwidth, screenheight = self.get_screen_size()
    self.center_window(self.root, width, height, screenwidth, screenheight)


  # 子窗口居中
  # 居中定位函数可能会被多个子窗口调用，因此引入window做参数，而不是self.root将其写死
  def center_child_window(self, window, width, height):
    screenwidth, screenheight = self.main_window_pos()
    self.center_window(window, width, height, screenwidth, screenheight)


  # 显示布局
  def show(self):
    self.top.pack(padx=12, pady=12)
    self.file_entry.pack(expand=tkinter.YES, side=tkinter.LEFT, fill=tkinter.Y, pady=1)
    # self.file_button.pack(side=tkinter.LEFT, fill=tkinter.BOTH, padx=8)
    self.parse_button.pack(expand=tkinter.YES, side=tkinter.LEFT, fill=tkinter.Y, padx=8, ipady=1)
    self.export_button.pack(expand=tkinter.YES, side=tkinter.LEFT, fill=tkinter.Y, ipady=1)

    self.middle.pack(padx=12, pady=0)
    self.result_list.pack(side=tkinter.LEFT, fill=tkinter.BOTH)
    # self.scroll_bar.pack(expand=tkinter.YES, fill=tkinter.Y)

    self.bottom.pack(fill=tkinter.BOTH, padx=12, pady=8)
    self.author_label.pack(side=tkinter.LEFT, fill=tkinter.Y)
    self.author_button.pack(side=tkinter.RIGHT, fill=tkinter.Y)

    self.root.mainloop()


  # 文件选择对话框
  def chooseFile(self):
    filename = tkinter.filedialog.askopenfilename(filetypes=[('Excel 97-2003 工作簿', ".xls")])
    self.file_entry.delete(0, tkinter.END)
    self.file_entry.insert(0, filename)
    self.filename = filename

  def chooseFolder(self):
    foldername = tkinter.filedialog.askdirectory()
    self.foldername = foldername


  def isFindKey(self, string):
    string = preTreat(string)
    for index, keyword in enumerate(self.keywords):
      if string == keyword:
        del self.keywords[index] # 关键词已摘取，将其从列表中删除，防止重复
        #print(string, '1', keyword)
        return True, keyword
    # for index, keyword in enumerate(self.keywords):
    #   if string.startswith(keyword) and Levenshtein.jaro(string, keyword) > 0.92:
    #     print(string, '1')
    #     del self.keywords[index]  # 关键词已摘取，将其从列表中删除，防止重复
    #     return True, keyword
    for index, keyword in enumerate(self.keywords):
      #print(string, keyword)
      if Levenshtein.jaro(string, keyword) > 0.92:
        #print(string, '2', Levenshtein.jaro(string, keyword), keyword)
        del self.keywords[index]  # 关键词已摘取，将其从列表中删除，防止重复
        return True, keyword
    return False, ''


  # 函数功能与countSheet一样，只不过为了在扫描并计算Excel数据的同时
  # 将结果渲染到GUI页面的List中，将函数进行局部修改定义为ParseExcel的成员函数
  def countSheetList(self, wb, sheet_name):
    sheet = wb.sheet_by_name(sheet_name)
    n_rows = sheet.nrows
    n_cols = sheet.ncols
    start_index = []
    end_index = []
    for i in range(n_rows):
      for j in range(n_cols):
        type_x_y, index_x_y = getCell(sheet, i, j)
        if type_x_y == biffh.XL_CELL_TEXT and not matchNumber(index_x_y):
          if 'company' not in self.results and (str(index_x_y).startswith('单位名称') or str(index_x_y).startswith('编制单位') or str(index_x_y).endswith('公司')):
            add2dict(self.results, 'company', str(index_x_y)[5:])
          if index_x_y == '期末余额' or index_x_y == '本期金额':
            end_index.append(j)
          if index_x_y == '年初余额' or index_x_y == '上期金额':
            start_index.append(j)

          flag, key = self.isFindKey(index_x_y)
          if(flag):
            tmp = countCell(sheet, end_index, i, j)
            # self.result_list.insert(tkinter.END, key)
            # self.result_list.insert(tkinter.END, tmp)
            # self.result_list.insert(tkinter.END, (key, tmp['start'], tmp['end'], tmp['change']))
            self.result_list.insert('', 'end', values=(key, tmp['start'], tmp['end'], tmp['change']))
            add2dict(self.results, key, tmp)


  def clear(self, tree):
    tmp = tree.get_children()
    for item in tmp:
      tree.delete(item)


  # 执行解析
  def conduct(self):
    self.chooseFile()
    self.keywords = ['非流动资产合计', '流动资产合计', '资产总计', '非流动负债合计', '流动负债合计', '负债总计', '负债合计', '实收资本', '所有者权益合计', '营业收入', '营业成本', '利润总额', '净利润']
    self.results = {}  # 初始化
    self.clear(self.result_list)
    if self.filename != '':
      wb = open_workbook(self.filename)
      sheet_list = wb.sheet_names()
      for sheet in sheet_list:
        if sheet.find('资产负债') != -1 or sheet.find('利润') != -1:
          self.countSheetList(wb, sheet)
    # else:
    #   # 定义子窗口
    #   self.child = tkinter.Toplevel()
    #   self.child.title('错误提示')
    #   self.center_child_window(self.child, 300, 90)
    #   err_msg = tkinter.Message(self.child, text = '错误：未选择文件！！！', width=150, padx=50, pady=30, anchor=tkinter.CENTER)
    #   err_msg.pack()
    #   self.child.mainloop()


  def export(self):
    if self.filename != '':
      rows_order = ['流动资产合计', '非流动资产合计', '资产总计', '流动负债合计', '非流动负债合计', '负债合计', '实收资本', '所有者权益合计', '营业收入', '营业成本', '利润总额', '净利润']
      self.chooseFolder()
      '''
      doc = Document('老板希望的导出效果.docx')
      table = doc.tables[0]
      for row in table.rows[1:]:
        key = row.cells[0].text
        row.cells[1].text = str(self.results[key]['end'])
        row.cells[2].text = str(self.results[key]['start'])
      para = doc.paragraphs
      for p in para:
        for r in p.runs:
          if r.underline:
            print(r.underline)
          else:
            print(r.text)
      doc.save('test.docx')
      '''
      doc = Document()

      doc.styles['Normal'].font.name = '宋体'
      doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

      doc.add_heading(self.results['company'], 0)

      table = doc.add_table(rows=1, cols=4, style='Table Grid')
      hdr_cells = table.rows[0].cells
      hdr_cells[0].text = '科目'
      hdr_cells[1].text = '年初（万元）'
      hdr_cells[2].text = '期末（万元）'
      hdr_cells[3].text = '变动（%）'
      for item in rows_order:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        # 鉴于数据的不完整性，可能会出现KeyError异常（item）导致程序退出，这里使用try来捕捉异常，做兼容性处理
        try:
          row_cells[1].text = str(self.results[item]['start'])
          row_cells[2].text = str(self.results[item]['end'])
          row_cells[3].text = str(self.results[item]['change'])
        except KeyError:
          row_cells[1].text = row_cells[2].text = row_cells[3].text = ''

      try:
        ts = self.results['资产总计']
      except KeyError:
        ts = {'start': ' ', 'end': ' ', 'change': ' '}
      try:
        td = self.results['负债合计']
      except KeyError:
        td = {'start': ' ', 'end': ' ', 'change': ''}
      try:
        oi = self.results['营业收入']
      except KeyError:
        oi = {'start': ' ', 'end': ' ', 'change': ' '}
      try:
        np = self.results['净利润']
      except KeyError:
        np = {'start': ' ', 'end': ' ', 'change': ' '}
      doc.add_paragraph('截至报告期末，项目公司总资产{0}万元，相对上期增长{1}%；总负债{2}万元，相对上期增长{3}%。收入和利润方面，报告期末公司营业收入{4}万元，相对上期增长{5}%；净利润{6}万元，相对上期增长{7}%。'.format(ts['end'], ts['change'], td['end'], td['change'], oi['end'], oi['change'], np['end'], np['change']))

      doc.add_paragraph('通过上述信息，公司各项财务指标（有/无）异常，报告期内（有/无）明显变化。总体来看，公司财务状况（）。')

      if self.foldername !='':
        doc.save(os.path.join(self.foldername, '{}.docx'.format(self.results['company'])))
      else:
        doc.save(os.path.join(os.getcwd(), '{}.docx'.format(self.results['company'])))
        tkinter.messagebox.showinfo('提示', '未选择文件存放位置，默认存放于当前文件夹！')
    else:
      tkinter.messagebox.showerror('错误', '未选择文件！')

  def broswer(self):
    webbrowser.open_new('https://github.com/Levance/FS-Parser')


def main():
  PE = ParseExcel()
  PE.show()
  #tkinter.mainloop()
  pass


if __name__ == '__main__':
  main()
